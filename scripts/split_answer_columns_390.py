# -*- coding: utf-8 -*-
"""
Tách cột \"Đáp án / Ghi chú\" thành 3 cột: Đáp án | Ghi chú | Căn cứ.
Đọc Word 4 cột đã biên tập, ghi Word/CSV/JSON 6 cột.
"""

from __future__ import annotations

import csv
import json
import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DIR = Path(r"g:\My Drive\AIProject_Data\Thi CCHN\Dau thau")
IN_DOCX = DIR / "NVCM-dau-thau-390-cau-bang.docx"
BACKUP_DOCX = DIR / "NVCM-dau-thau-390-cau-bang-4cot-backup.docx"
OUT_DOCX = DIR / "NVCM-dau-thau-390-cau-bang.docx"
OUT_DOCX_ALT = DIR / "NVCM-dau-thau-390-cau-bang-6cot.docx"
OUT_CSV = DIR / "NVCM-dau-thau-390-cau.csv"
OUT_JSON = DIR / "NVCM-dau-thau-390-cau.json"

GREEN = RGBColor(0x00, 0x80, 0x00)
RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

TAG_RE = re.compile(r"\((?:bổ sung|mới)\)", re.IGNORECASE)
DUP_ANSWER_RE = re.compile(
    r"^[A-D](?:\s*\|\s*[A-D])*(?:\s*\((?:bổ sung|mới)\))?$",
    re.IGNORECASE,
)
CROSS_NOTE_RE = re.compile(r"^(?:\d*[a-zA-Z]*c\s*=\s*[A-D]|4a\s*=\s*[A-D])$", re.IGNORECASE)
LEADING_ANSWER_RE = re.compile(
    r"^([A-D](?:\s*\|\s*[A-D])*)(?:\s*(\((?:bổ sung|mới)\)))?\s*(.*)$",
    re.IGNORECASE | re.DOTALL,
)


def norm_answer_expr(s: str) -> str:
    s = (s or "").upper().replace(" ", "")
    return re.sub(r"\|", " | ", s)


def set_run_font(run, *, size: int = 11, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_text(
    cell,
    text: str,
    *,
    size: int = 11,
    bold: bool = False,
    align=None,
    color: RGBColor | None = None,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def set_multiline_cell(cell, text: str, *, size: int = 10) -> None:
    cell.text = ""
    lines = (text or "").split("\n")
    first = True
    for line in lines:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(line)
        set_run_font(run, size=size, bold=False)


def shade_header(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), "D9E2F3")
    shd.set(qn("w:val"), "clear")


def set_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "666666")
        borders.append(el)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def cell_answer_color(cell) -> RGBColor:
    """Giữ màu chữ đáp án (xanh bổ sung / đỏ khác bản cũ / đen)."""
    for para in cell.paragraphs:
        for run in para.runs:
            if not run.text.strip():
                continue
            rgb = run.font.color.rgb if run.font.color and run.font.color.rgb else None
            if rgb is None:
                return BLACK
            s = str(rgb).upper()
            if s in {"008000", "00B050", "00FF00"} or (
                rgb[1] >= 0x70 and rgb[0] <= 0x40 and rgb[2] <= 0x40
            ):
                return GREEN
            if s in {"FF0000", "C00000", "FF0000"} or (
                rgb[0] >= 0xB0 and rgb[1] <= 0x40 and rgb[2] <= 0x40
            ):
                return RED
            return BLACK
    return BLACK


def unwrap_parens(s: str) -> str:
    s = s.strip()
    if len(s) >= 2 and s[0] == "(" and s[-1] == ")":
        depth = 0
        for i, ch in enumerate(s):
            if ch == "(":
                depth += 1
            elif ch == ")":
                depth -= 1
                if depth == 0 and i != len(s) - 1:
                    return s
        return s[1:-1].strip()
    return s


def dedupe_notes(parts: list[str]) -> str:
    seen: set[str] = set()
    tags: list[str] = []
    others: list[str] = []
    for p in parts:
        p = re.sub(r"\s+", " ", p).strip(" ·;")
        if not p:
            continue
        key = p.lower()
        if key in seen:
            continue
        seen.add(key)
        if TAG_RE.fullmatch(p):
            tags.append(p)
        else:
            others.append(p)
    return "; ".join(tags + others)


def split_answer_cell(raw: str) -> tuple[str, str, str]:
    """
    Trả về (đáp_án, ghi_chú, căn_cứ).

    Ví dụ:
      A — (Khoản 1, Điều 2, ...) · 2c=D
        -> A | 2c=D | Khoản 1, Điều 2, ...
      C (bổ sung) — C — C — 2c=D · 2c=D
        -> C | (bổ sung); 2c=D | (trống)
      B (Khoản 4, Điều 27, Thông tư 79)
        -> B | | Khoản 4, Điều 27, Thông tư 79
    """
    text = (raw or "").replace("\xa0", " ").strip()
    if not text:
        return "", "", ""

    note_parts: list[str] = []

    # Tách ghi chú sau dấu · (ví dụ 2c=D)
    if "·" in text:
        chunks = [c.strip() for c in re.split(r"\s*·\s*", text) if c.strip()]
        text = chunks[0] if chunks else ""
        note_parts.extend(chunks[1:])

    # Không có gạch ngang: "D" hoặc "B (căn cứ...)"
    if "—" not in text and "–" not in text:
        m = LEADING_ANSWER_RE.match(text)
        if m:
            answer = norm_answer_expr(m.group(1))
            if m.group(2):
                note_parts.append(m.group(2))
            rest = (m.group(3) or "").strip()
            return answer, dedupe_notes(note_parts), unwrap_parens(rest)
        return text, dedupe_notes(note_parts), ""

    # Tách theo — / – (có hoặc không khoảng trắng)
    parts = [p.strip() for p in re.split(r"\s*[—–]\s*", text) if p.strip()]
    left = parts[0] if parts else ""
    rest_parts = parts[1:]

    m = LEADING_ANSWER_RE.match(left)
    if m:
        answer = norm_answer_expr(m.group(1))
        if m.group(2):
            note_parts.append(m.group(2))
        leftover = (m.group(3) or "").strip()
        if leftover:
            # Trường hợp hiếm: căn cứ dính luôn ở phần trái
            if leftover.startswith("(") or re.search(
                r"(?i)(khoản|điều|điểm|mục|luật|nghị định|thông tư)", leftover
            ):
                rest_parts.insert(0, leftover)
            else:
                note_parts.append(leftover)
    else:
        answer = left

    can_cu_parts: list[str] = []
    for p in rest_parts:
        if CROSS_NOTE_RE.fullmatch(p.replace(" ", "")):
            note_parts.append(re.sub(r"\s+", "", p))
            continue
        if DUP_ANSWER_RE.fullmatch(p):
            expr = norm_answer_expr(TAG_RE.sub("", p))
            # Chỉ giữ khi khác đáp án chính (đối chiếu / xung đột)
            if expr and expr != answer:
                note_parts.append(expr)
            continue
        can_cu_parts.append(unwrap_parens(p))

    can_cu = "; ".join(x for x in can_cu_parts if x)
    return answer, dedupe_notes(note_parts), can_cu


def load_rows(path: Path) -> list[dict]:
    doc = Document(str(path))
    if not doc.tables:
        raise SystemExit(f"No table in {path}")
    table = doc.tables[0]
    rows: list[dict] = []
    for row in table.rows[1:]:
        stt = row.cells[0].text.strip()
        prompt = row.cells[1].text.strip()
        options = row.cells[2].text.strip()
        raw_ans = row.cells[3].text.strip()
        color = cell_answer_color(row.cells[3])
        answer, note, basis = split_answer_cell(raw_ans)
        rows.append(
            {
                "stt": int(stt) if stt.isdigit() else stt,
                "prompt": prompt,
                "options_text": options,
                "answer": answer,
                "note": note,
                "basis": basis,
                "raw": raw_ans,
                "color": color,
            }
        )
    return rows


def parse_options_text(text: str) -> dict[str, str]:
    opts = {"A": "", "B": "", "C": "", "D": ""}
    parts = re.split(r"(?m)^([A-D])\.\s*", text)
    # parts: ['', 'A', 'text', 'B', 'text', ...]
    i = 1
    while i + 1 < len(parts):
        key = parts[i].upper()
        val = parts[i + 1].strip()
        if key in opts:
            opts[key] = val
        i += 2
    return opts


def write_docx(rows: list[dict], path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run(
        "Ngân hàng 390 câu NVCM đấu thầu — đã biên tập đáp án (tách cột Đáp án / Ghi chú / Căn cứ)"
    )
    set_run_font(r, size=13, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(8)
    r2 = sub.add_run(
        "Chú thích cột Đáp án: đỏ = đáp án mới khác bản cũ · xanh = bổ sung đáp án còn thiếu · đen = trùng khớp / giữ nguyên"
    )
    set_run_font(r2, size=9, bold=False)

    headers = (
        "STT",
        "Nội dung câu hỏi",
        "Phương án trả lời\n(A, B, C, D)",
        "Đáp án",
        "Ghi chú",
        "Căn cứ",
    )
    table = doc.add_table(rows=1 + len(rows), cols=6)
    table.style = "Table Grid"
    table.autofit = False
    set_table_borders(table)
    # Tổng ~27.3 cm nội dung (khổ ngang A4 trừ lề)
    widths = (Cm(1.0), Cm(7.2), Cm(10.0), Cm(1.6), Cm(2.8), Cm(4.7))
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w

    for i, h in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[i],
            h,
            size=10,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_header(table.rows[0].cells[i])

    for qi, q in enumerate(rows):
        row = table.rows[qi + 1]
        set_cell_text(
            row.cells[0],
            str(q["stt"]),
            size=10,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_cell_text(row.cells[1], q["prompt"], size=9)
        set_multiline_cell(row.cells[2], q["options_text"], size=9)
        set_cell_text(
            row.cells[3],
            q["answer"],
            size=10,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            color=q["color"],
        )
        set_cell_text(row.cells[4], q["note"], size=8)
        set_cell_text(row.cells[5], q["basis"], size=8)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def write_csv(rows: list[dict], path: Path) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(
            [
                "stt",
                "prompt",
                "option_a",
                "option_b",
                "option_c",
                "option_d",
                "answer",
                "note",
                "basis",
            ]
        )
        for q in rows:
            opts = parse_options_text(q["options_text"])
            w.writerow(
                [
                    q["stt"],
                    q["prompt"],
                    opts["A"],
                    opts["B"],
                    opts["C"],
                    opts["D"],
                    q["answer"],
                    q["note"],
                    q["basis"],
                ]
            )


def write_json(rows: list[dict], path: Path) -> None:
    out = []
    for q in rows:
        opts = parse_options_text(q["options_text"])
        out.append(
            {
                "stt": q["stt"],
                "prompt": q["prompt"],
                "options": opts,
                "answer": q["answer"],
                "note": q["note"],
                "basis": q["basis"],
            }
        )
    path.write_text(json.dumps(out, ensure_ascii=False, indent=2), encoding="utf-8")


def save_docx_with_fallback(rows: list[dict]) -> Path:
    try:
        write_docx(rows, OUT_DOCX)
        return OUT_DOCX
    except PermissionError:
        write_docx(rows, OUT_DOCX_ALT)
        print(
            f"File gốc đang mở trong Word — đã ghi tạm: {OUT_DOCX_ALT.name}\n"
            f"Hãy đóng Word rồi đổi tên file này thành {OUT_DOCX.name} (hoặc chạy lại script)."
        )
        return OUT_DOCX_ALT


def main() -> None:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")

    if not IN_DOCX.exists() and not BACKUP_DOCX.exists():
        raise SystemExit(f"Missing: {IN_DOCX}")

    # Ưu tiên đọc bản 4 cột đã backup (nếu file chính đã bị khóa / đã tách)
    source = IN_DOCX if IN_DOCX.exists() else BACKUP_DOCX
    probe = Document(str(source))
    ncols = len(probe.tables[0].rows[0].cells) if probe.tables else 0
    if ncols != 4:
        if BACKUP_DOCX.exists():
            source = BACKUP_DOCX
            print(f"File chính đã {ncols} cột — đọc từ backup 4 cột")
        else:
            raise SystemExit(f"Expected 4-column table, got {ncols} cols and no backup")
    elif not BACKUP_DOCX.exists():
        shutil.copy2(source, BACKUP_DOCX)
        print(f"Backup: {BACKUP_DOCX.name}")
    else:
        print(f"Backup already exists: {BACKUP_DOCX.name}")

    rows = load_rows(source)
    print(f"Loaded {len(rows)} rows from {source.name}")

    # Thống kê nhanh
    empty_ans = sum(1 for r in rows if not r["answer"])
    with_note = sum(1 for r in rows if r["note"])
    with_basis = sum(1 for r in rows if r["basis"])
    print(f"Đáp án trống: {empty_ans} | có Ghi chú: {with_note} | có Căn cứ: {with_basis}")

    # In vài mẫu kiểm tra
    for idx in (0, 27, 28, 64, 260, 324):
        if idx < len(rows):
            r = rows[idx]
            print(
                f"#{r['stt']}: ans={r['answer']!r} | note={r['note']!r} | basis={r['basis'][:80]!r}"
            )
            print(f"     raw={r['raw'][:100]!r}")

    out_path = save_docx_with_fallback(rows)
    write_csv(rows, OUT_CSV)
    write_json(rows, OUT_JSON)
    print(f"Wrote: {out_path.name}")
    print(f"Wrote: {OUT_CSV.name}")
    print(f"Wrote: {OUT_JSON.name}")


if __name__ == "__main__":
    main()
