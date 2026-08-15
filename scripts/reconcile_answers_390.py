# -*- coding: utf-8 -*-
"""
Đối chiếu đáp án ngân hàng 390 với:
  - 4a (PDF 390/395 câu có cột Đáp án)
  - 2c (PDF 270 câu — đáp án in đậm)
Cập nhật Word: đáp án mới đỏ nếu khác; bổ sung xanh nếu thiếu.
"""

from __future__ import annotations

import csv
import json
import re
import sys
import unicodedata
from pathlib import Path

import fitz
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DIR = Path(r"g:\My Drive\AIProject_Data\Thi CCHN\Dau thau")
BASE = DIR / (
    "CTy Tu van CDH - On thi dau thau 2026 - 390 cau-20260815T085047Z-1-001"
)
F4A = BASE / "4a - trắc nghiệm đấu thầu nâng cao (cập nhập đến 22.01.26) - 390 câu.pdf"
F2C = BASE / "2c - DAP AN 270 CAU CCNVCM.pdf"
JSON_IN = DIR / "NVCM-dau-thau-390-cau.json"
OUT_DOCX = DIR / "NVCM-dau-thau-390-cau-bang.docx"
OUT_CSV = DIR / "NVCM-dau-thau-390-cau.csv"
OUT_JSON = DIR / "NVCM-dau-thau-390-cau.json"
OUT_REPORT = DIR / "NVCM-dau-thau-390-doi-chieu-dap-an.txt"

RED = RGBColor(0xC0, 0x00, 0x00)
GREEN = RGBColor(0x00, 0x80, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)


def norm(s: str) -> str:
    s = unicodedata.normalize("NFC", s or "")
    s = s.lower()
    s = re.sub(r"\s+", " ", s)
    s = re.sub(r"[“”\"'`]", "", s)
    return s.strip()


def prompt_key(s: str, n: int = 90) -> str:
    s = norm(s)
    s = re.sub(r"[^\w\sÀ-ỹ]", "", s, flags=re.UNICODE)
    return s[:n]


def letter(s: str | None) -> str:
    if not s:
        return ""
    s = str(s).strip().upper()
    m = re.search(r"[ABCD]", s)
    return m.group(0) if m else ""


def parse_4a(path: Path) -> dict[int, dict]:
    doc = fitz.open(str(path))
    out: dict[int, dict] = {}
    for page in doc:
        tabs = page.find_tables()
        if not tabs or not tabs.tables:
            continue
        for table in tabs.tables:
            for row in table.extract() or []:
                if not row or not str(row[0] or "").strip().isdigit():
                    continue
                qid = int(str(row[0]).strip())
                prompt = re.sub(r"\s+", " ", str(row[1] or "")).strip()
                ans = letter(str(row[6] if len(row) > 6 else row[-1] or ""))
                if not ans:
                    continue
                out[qid] = {"stt": qid, "prompt": prompt, "answer": ans}
    return out


def parse_2c(path: Path) -> dict[int, dict]:
    """Đáp án = dòng A/B/C/D in đậm (thường kèm dẫn chứng trong ngoặc)."""
    doc = fitz.open(str(path))
    answers: dict[int, dict] = {}
    opt_re = re.compile(r"^([A-D])\.\s*(.*)$")
    stt_re = re.compile(r"^(\d{1,3})$")

    for page in doc:
        # Build line records from spans
        d = page.get_text("dict")
        page_lines: list[tuple[float, float, str, bool]] = []
        for b in d["blocks"]:
            if b.get("type") != 0:
                continue
            for line in b.get("lines", []):
                spans = [s for s in line.get("spans", []) if s["text"].strip()]
                if not spans:
                    continue
                text = "".join(s["text"] for s in spans).strip()
                bold_len = sum(
                    len(s["text"]) for s in spans if (s.get("flags", 0) & 16)
                )
                total_len = sum(len(s["text"]) for s in spans) or 1
                bold = bold_len / total_len >= 0.4
                page_lines.append((line["bbox"][1], line["bbox"][0], text, bold))
        page_lines.sort(key=lambda t: (t[0], t[1]))

        current: int | None = None
        prompt_parts: list[str] = []
        for y0, x0, text, bold in page_lines:
            if text in {"STT", "Nội dung", "Phương án trả lời"}:
                continue
            if text.startswith("ĐÁP ÁN") or text.startswith("CẤP CCN"):
                continue
            if "Tham khảo đáp án" in text or text.startswith("ý kiến"):
                continue

            m_stt = stt_re.match(text)
            # STT thường cột trái; nới điều kiện x
            if m_stt and x0 < 120:
                n = int(m_stt.group(1))
                if 1 <= n <= 300:
                    current = n
                    prompt_parts = []
                    answers[n] = {"stt": n, "prompt": "", "answer": "", "cite": ""}
                continue
            if current is None:
                continue

            m_opt = opt_re.match(text)
            if m_opt:
                mode = m_opt.group(1)
                # Chỉ lấy đáp án đậm; nếu đã có rồi thì bỏ qua (tránh dính câu sau)
                if bold and not answers[current]["answer"]:
                    answers[current]["answer"] = mode
                    cites = re.findall(r"\(([^)]+)\)", text)
                    if cites:
                        answers[current]["cite"] = cites[-1]
                continue

            # Dòng tiếp theo của đáp án đậm: dẫn chứng
            if (
                bold
                and answers[current]["answer"]
                and not answers[current]["cite"]
            ):
                cites = re.findall(r"\(([^)]+)\)", text)
                if cites:
                    answers[current]["cite"] = cites[-1]
                continue

            if x0 < 280 and not opt_re.match(text):
                prompt_parts.append(text)
                answers[current]["prompt"] = " ".join(prompt_parts)

    for q in answers.values():
        q["prompt"] = re.sub(r"\s+", " ", q["prompt"]).strip()
    return answers


def build_prompt_index(items: dict[int, dict]) -> dict[str, int]:
    idx: dict[str, int] = {}
    for qid, q in items.items():
        k = prompt_key(q["prompt"])
        if k and k not in idx:
            idx[k] = qid
    return idx


def find_match(prompt: str, by_id: dict[int, dict], pindex: dict[str, int]) -> dict | None:
    from difflib import SequenceMatcher

    k = prompt_key(prompt, 120)
    if not k:
        return None

    # Khớp exact key
    if k in pindex:
        # với câu ngắn: chỉ nhận nếu duy nhất trong bộ
        if len(k) < 45:
            twins = [qid for pk, qid in pindex.items() if pk == k or pk.startswith(k) or k.startswith(pk)]
            # cũng đếm số câu 4a có cùng đoạn đầu ngắn
            short = k[:40]
            near = [qid for pk, qid in pindex.items() if pk[:40] == short]
            if len(set(near)) > 1:
                return None
        return by_id[pindex[k]]

    if len(k) < 35:
        # quá ngắn và không exact → bỏ, tránh khớp nhầm
        return None

    best = None
    best_r = 0.0
    k80 = k[:80]
    candidates = []
    for pk, qid in pindex.items():
        if not pk:
            continue
        if pk.startswith(k80) or k80.startswith(pk[: min(80, len(pk))]):
            r = SequenceMatcher(None, k[:100], pk[:100]).ratio()
            candidates.append((r, qid))
            if r > best_r:
                best_r, best = r, qid
            continue
        r = SequenceMatcher(None, k[:100], pk[:100]).ratio()
        if r >= 0.78:
            candidates.append((r, qid))
        if r > best_r:
            best_r, best = r, qid

    if best is None or best_r < 0.82:
        return None
    # nếu có từ 2 ứng viên gần bằng nhau → bỏ
    strong = [c for c in candidates if c[0] >= best_r - 0.03]
    if len({c[1] for c in strong}) > 1:
        return None
    return by_id[best]


def set_run_font(run, name="Times New Roman", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_plain(cell, text, *, size=11, bold=False, align=None, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_answer_cell(cell, old: str, new: str, status: str, extra_note: str = "") -> None:
    """status: same | conflict | filled | empty"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15

    def add(text, *, bold=False, color=None):
        run = p.add_run(text)
        set_run_font(run, size=10, bold=bold, color=color)

    if status == "same":
        add(old or new, bold=True, color=BLACK)
        if extra_note:
            add(f" — {extra_note}", color=BLACK)
    elif status == "conflict":
        # đáp án cũ | đáp án mới (đỏ)
        add(old or "?", bold=True, color=BLACK)
        add(" | ", color=BLACK)
        add(new, bold=True, color=RED)
        add(" (mới)", color=RED)
        if extra_note:
            add(f" — {extra_note}", color=BLACK)
    elif status == "filled":
        add(new, bold=True, color=GREEN)
        add(" (bổ sung)", color=GREEN)
        if extra_note:
            add(f" — {extra_note}", color=BLACK)
    else:
        if extra_note:
            add(extra_note, color=BLACK)
        else:
            add("", color=BLACK)


def add_options_cell(cell, options: list[str]):
    cell.text = ""
    first = True
    for key, text in zip("ABCD", options):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(f"{key}. {text}")
        set_run_font(run, size=10)


def shade_header(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), "D9E2F3")
    shd.set(qn("w:val"), "clear")


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "666666")
        borders.append(el)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def write_docx(bank: list[dict], path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Ngân hàng 390 câu NVCM đấu thầu — đã đối chiếu đáp án (4a + 2c)")
    set_run_font(r, size=13, bold=True)

    legend = doc.add_paragraph()
    legend.alignment = WD_ALIGN_PARAGRAPH.CENTER
    legend.paragraph_format.space_after = Pt(10)
    r1 = legend.add_run("Chú thích cột Đáp án/Ghi chú: ")
    set_run_font(r1, size=10)
    r2 = legend.add_run("đỏ = đáp án mới khác bản cũ")
    set_run_font(r2, size=10, bold=True, color=RED)
    r3 = legend.add_run(" · ")
    set_run_font(r3, size=10)
    r4 = legend.add_run("xanh = bổ sung đáp án còn thiếu")
    set_run_font(r4, size=10, bold=True, color=GREEN)
    r5 = legend.add_run(" · đen = trùng khớp / giữ nguyên")
    set_run_font(r5, size=10)

    headers = (
        "STT",
        "Nội dung câu hỏi",
        "Phương án trả lời\n(A, B, C, D)",
        "Đáp án / Ghi chú",
    )
    table = doc.add_table(rows=1 + len(bank), cols=4)
    table.style = "Table Grid"
    table.autofit = False
    set_table_borders(table)
    widths = (Cm(1.2), Cm(8.3), Cm(13.0), Cm(4.2))
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w
    for i, h in enumerate(headers):
        set_cell_plain(
            table.rows[0].cells[i],
            h,
            size=11,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_header(table.rows[0].cells[i])

    for qi, q in enumerate(bank):
        row = table.rows[qi + 1]
        set_cell_plain(
            row.cells[0],
            str(q["stt"]),
            size=11,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_cell_plain(row.cells[1], q["prompt"], size=10)
        add_options_cell(row.cells[2], q["options"])
        add_answer_cell(
            row.cells[3],
            q.get("answer_old") or "",
            q.get("answer") or "",
            q.get("answer_status") or "empty",
            q.get("note_extra") or "",
        )

    doc.save(str(path))


def main() -> None:
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

    print("Parsing 4a...")
    a4 = parse_4a(F4A)
    print(f"  4a: {len(a4)} answers")
    print("Parsing 2c...")
    a2 = parse_2c(F2C)
    print(f"  2c: {len(a2)} answers; with letter: {sum(1 for q in a2.values() if q['answer'])}")

    idx4 = build_prompt_index(a4)
    idx2 = build_prompt_index(a2)

    bank = json.loads(JSON_IN.read_text(encoding="utf-8"))
    stats = {"same": 0, "conflict": 0, "filled": 0, "empty": 0, "matched4a": 0, "matched2c": 0}
    report_lines = []

    for item in bank:
        # Mốc đối chiếu = đáp án gốc trước lần rà soát (nếu có)
        if "answer_old" in item:
            old = letter(item.get("answer_old"))
        else:
            old = letter(item.get("answer"))
        src = item.get("source")
        src_stt = item.get("source_stt")
        # Ưu tiên khớp theo nội dung câu hỏi; fallback theo STT nguồn
        ref4 = find_match(item["prompt"], a4, idx4)
        if (
            not ref4
            and src == "DA340"
            and isinstance(src_stt, int)
            and src_stt in a4
        ):
            ref4 = a4[src_stt]

        ref2 = find_match(item["prompt"], a2, idx2)
        if (
            not ref2
            and src == "DA340"
            and isinstance(src_stt, int)
            and src_stt in a2
        ):
            ref2 = a2[src_stt]

        ans4 = letter(ref4["answer"]) if ref4 else ""
        ans2 = letter(ref2["answer"]) if ref2 else ""
        if ref4:
            stats["matched4a"] += 1
        if ref2:
            stats["matched2c"] += 1

        # Nguồn ưu tiên: 4a (bộ cập nhật), fallback 2c
        new = ans4 or ans2

        cite = ""
        if ref2 and ref2.get("cite"):
            cite = ref2["cite"]
        old_note = item.get("note") or ""
        old_cite = re.sub(r"^[ABCD]\s*—\s*", "", old_note).strip()
        old_cite = re.sub(
            r"\s*\(mới\)|\s*\(bổ sung\)|·\s*2c=[ABCD]", "", old_cite
        ).strip(" —")
        # bỏ ghi chú kỹ thuật cũ
        if old_cite.startswith("Thiếu phương án") or old_cite.startswith("Chưa có đáp án"):
            old_cite = ""

        extra = cite or old_cite
        if ans4 and ans2 and ans4 != ans2:
            extra = (extra + " · " if extra else "") + f"2c={ans2}"

        item["answer_old"] = old
        item["answer_4a"] = ans4 or None
        item["answer_2c"] = ans2 or None
        item["ref_4a_id"] = ref4["stt"] if ref4 else None
        item["ref_2c_id"] = ref2["stt"] if ref2 else None

        if old and new:
            if old == new:
                item["answer"] = old
                item["answer_status"] = "same"
                item["note"] = old + (f" — {extra}" if extra else "")
                item["note_extra"] = extra
                stats["same"] += 1
            else:
                item["answer"] = new
                item["answer_status"] = "conflict"
                item["note"] = f"{old} | {new} (mới)" + (f" — {extra}" if extra else "")
                item["note_extra"] = extra
                stats["conflict"] += 1
                report_lines.append(
                    f"CONFLICT STT{item['stt']}: old={old} new={new} "
                    f"4a_id={item['ref_4a_id']} 2c_id={item['ref_2c_id']} | {item['prompt'][:80]}"
                )
        elif not old and new:
            item["answer"] = new
            item["answer_status"] = "filled"
            item["note"] = f"{new} (bổ sung)" + (f" — {extra}" if extra else "")
            item["note_extra"] = extra
            stats["filled"] += 1
            report_lines.append(
                f"FILLED STT{item['stt']}: {new} 4a_id={item['ref_4a_id']} | {item['prompt'][:80]}"
            )
        elif old and not new:
            item["answer"] = old
            item["answer_status"] = "same"
            item["note"] = old + (f" — {extra}" if extra else "")
            item["note_extra"] = extra
            stats["same"] += 1
        else:
            item["answer"] = None
            item["answer_status"] = "empty"
            item["note"] = ""
            item["note_extra"] = ""
            stats["empty"] += 1
            report_lines.append(f"EMPTY STT{item['stt']}: {item['prompt'][:100]}")

    print("Stats:", stats)
    OUT_REPORT.write_text(
        "Thống kê đối chiếu đáp án\n"
        + json.dumps(stats, ensure_ascii=False, indent=2)
        + "\n\n"
        + "\n".join(report_lines),
        encoding="utf-8",
    )

    # write outputs
    OUT_JSON.write_text(json.dumps(bank, ensure_ascii=False, indent=2), encoding="utf-8")
    with OUT_CSV.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(
            [
                "stt",
                "prompt",
                "option_a",
                "option_b",
                "option_c",
                "option_d",
                "answer",
                "answer_old",
                "answer_4a",
                "answer_2c",
                "answer_status",
                "note",
                "source",
                "source_stt",
                "ref_4a_id",
                "ref_2c_id",
            ]
        )
        for q in bank:
            w.writerow(
                [
                    q["stt"],
                    q["prompt"],
                    q["options"][0],
                    q["options"][1],
                    q["options"][2],
                    q["options"][3],
                    q.get("answer") or "",
                    q.get("answer_old") or "",
                    q.get("answer_4a") or "",
                    q.get("answer_2c") or "",
                    q.get("answer_status") or "",
                    q.get("note") or "",
                    q.get("source") or "",
                    q.get("source_stt") or "",
                    q.get("ref_4a_id") or "",
                    q.get("ref_2c_id") or "",
                ]
            )

    print("Writing Word...")
    write_docx(bank, OUT_DOCX)
    print("Wrote", OUT_DOCX)
    print("Wrote", OUT_CSV)
    print("Wrote", OUT_JSON)
    print("Wrote", OUT_REPORT)


if __name__ == "__main__":
    main()
