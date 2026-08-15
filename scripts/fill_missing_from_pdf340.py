# -*- coding: utf-8 -*-
"""Điền phương án thiếu trong bộ 390 từ PDF Thông báo 1952 (340 câu)."""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

import fitz
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

DIR = Path(r"g:\My Drive\AIProject_Data\Thi CCHN\Dau thau")
PDF_340 = DIR / "Thông báo 1952-340 câu hỏi thi NVCM Đấu thầu (thay thế 1813).pdf"
JSON_IN = DIR / "NVCM-dau-thau-390-cau.json"
OUT_DOCX = DIR / "NVCM-dau-thau-390-cau-bang.docx"
OUT_CSV = DIR / "NVCM-dau-thau-390-cau.csv"
OUT_JSON = DIR / "NVCM-dau-thau-390-cau.json"

SKIP = {"STT", "Nội dung", "Phương án trả lời"}
OPTION_START = re.compile(r"^([A-D])\s*\.\s*(.*)$", re.S)
STT_ONLY = re.compile(r"^(\d{1,3})$")
TITLE_SKIP = re.compile(
    r"^(Thông báo|Bổ sung|Ngân hàng|Cục Quản|phục vụ)", re.I
)


def clean(s: str) -> str:
    s = re.sub(r"[ \t]+", " ", (s or "").replace("\xa0", " "))
    s = re.sub(r"\s+([,.;:!?])", r"\1", s)
    return s.strip()


def extract_lines(pdf: fitz.Document) -> list[str]:
    lines: list[str] = []
    for page in pdf:
        for line in page.get_text("text").splitlines():
            t = line.strip()
            if t:
                lines.append(t)
    return lines


def parse_pdf_questions(lines: list[str]) -> dict[int, dict]:
    questions: dict[int, dict] = {}
    current: dict | None = None
    mode: str | None = None

    def append(bucket: str, text: str) -> None:
        assert current is not None
        if bucket == "prompt":
            prev = current["prompt"]
            joiner = "" if not prev or prev.endswith((" ", "-", "/")) else " "
            current["prompt"] = prev + joiner + text
        else:
            prev = current["options"][bucket]
            joiner = "" if not prev or prev.endswith((" ", "-", "/")) else " "
            current["options"][bucket] = prev + joiner + text

    for text in lines:
        if text in SKIP or TITLE_SKIP.search(text):
            continue
        stt_m = STT_ONLY.match(text)
        if stt_m:
            n = int(stt_m.group(1))
            if 1 <= n <= 340:
                current = {
                    "stt": n,
                    "prompt": "",
                    "options": {"A": "", "B": "", "C": "", "D": ""},
                }
                questions[n] = current
                mode = "prompt"
                continue
        if current is None:
            continue
        opt_m = OPTION_START.match(text)
        if opt_m:
            mode = opt_m.group(1)
            rest = opt_m.group(2).strip()
            if rest:
                append(mode, rest)
            continue
        if mode:
            append(mode, text)

    for q in questions.values():
        q["prompt"] = clean(q["prompt"])
        for k in "ABCD":
            q["options"][k] = clean(q["options"][k])
    return questions


def set_run_font(run, name="Times New Roman", size=11, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_text(cell, text, *, size=11, bold=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_options_cell(cell, options, *, size=10):
    cell.text = ""
    first = True
    for key in ("A", "B", "C", "D"):
        text = f"{key}. {options[key]}"
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        set_run_font(run, size=size)


def shade_header(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), "D9E2F3")
    shd.set(qn("w:val"), "clear")


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "666666")
        borders.append(el)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def write_docx(questions: list[dict], path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "Ngân hàng câu hỏi NVCM đấu thầu — 390 câu (340 + 50 bổ sung TB 1952 lần 2)"
    )
    set_run_font(r, size=13, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(10)
    r2 = sub.add_run(
        "STT 1–340: theo PDF Thông báo 1952 (thay thế 1813). "
        "STT 341–390: bổ sung lần 2 — chưa có đáp án thì để trống."
    )
    set_run_font(r2, size=10)

    headers = (
        "STT",
        "Nội dung câu hỏi",
        "Phương án trả lời\n(Liệt kê đầy đủ A, B, C, D)",
        "Đáp án / Ghi chú",
    )
    table = doc.add_table(rows=1 + len(questions), cols=4)
    table.style = "Table Grid"
    table.autofit = False
    set_table_borders(table)
    widths = (Cm(1.2), Cm(8.3), Cm(13.2), Cm(4.0))
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w
    for i, h in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[i],
            h,
            size=11,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_header(table.rows[0].cells[i])

    for qi, q in enumerate(questions):
        row = table.rows[qi + 1]
        opts = {
            "A": q["options"][0],
            "B": q["options"][1],
            "C": q["options"][2],
            "D": q["options"][3],
        }
        set_cell_text(
            row.cells[0],
            str(q["stt"]),
            size=11,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_cell_text(row.cells[1], q["prompt"], size=10)
        add_options_cell(row.cells[2], opts, size=10)
        set_cell_text(row.cells[3], q.get("note") or "", size=10)

    doc.save(str(path))


def write_csv(questions: list[dict], path: Path) -> None:
    import csv

    with path.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(
            [
                "stt",
                "prompt",
                "option_a",
                "option_b",
                "option_c",
                "option_d",
                "answer",
                "note",
                "source",
                "source_stt",
            ]
        )
        for q in questions:
            w.writerow(
                [
                    q["stt"],
                    q["prompt"],
                    q["options"][0],
                    q["options"][1],
                    q["options"][2],
                    q["options"][3],
                    q.get("answer") or "",
                    q.get("note") or "",
                    q.get("source") or "",
                    q.get("source_stt") or "",
                ]
            )


def main() -> None:
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

    pdf = fitz.open(str(PDF_340))
    pdf_qs = parse_pdf_questions(extract_lines(pdf))
    print(f"Parsed PDF: {len(pdf_qs)} questions")
    if len(pdf_qs) != 340:
        missing = [i for i in range(1, 341) if i not in pdf_qs]
        print("Missing STTs in PDF parse:", missing[:30], f"... total {len(missing)}")

    incomplete_pdf = []
    for n, q in sorted(pdf_qs.items()):
        for k in "ABCD":
            if not q["options"][k]:
                incomplete_pdf.append(f"{n}{k}")
    if incomplete_pdf:
        print("PDF still incomplete:", incomplete_pdf)
    else:
        print("PDF parse: đủ A–D cho mọi câu")

    bank = json.loads(JSON_IN.read_text(encoding="utf-8"))
    filled = []
    still_bad = []

    for item in bank:
        src = item.get("source")
        src_stt = item.get("source_stt")
        if src == "DA340" and isinstance(src_stt, int) and src_stt in pdf_qs:
            pdf_q = pdf_qs[src_stt]
            opts = item["options"]
            changed = False
            for i, key in enumerate("ABCD"):
                pdf_opt = pdf_q["options"][key]
                if (not opts[i] or not str(opts[i]).strip()) and pdf_opt:
                    opts[i] = pdf_opt
                    changed = True
                    filled.append(f"STT{item['stt']}(src{src_stt}){key}")
                # nếu PDF đầy đủ hơn / sửa prompt bị dính A
            # luôn đồng bộ lại 4 phương án từ PDF cho các câu từng thiếu
            if any(not (opts[i] or "").strip() for i in range(4)) or item["stt"] in {
                95,
                294,
                297,
                298,
                319,
                321,
            }:
                for i, key in enumerate("ABCD"):
                    if pdf_q["options"][key]:
                        if opts[i] != pdf_q["options"][key]:
                            changed = True
                        opts[i] = pdf_q["options"][key]
                # prompt: nếu PDF có và nguồn bị cắt/dính
                if pdf_q["prompt"] and (
                    not item["prompt"]
                    or item["prompt"].endswith("A")
                    or len(pdf_q["prompt"]) > len(item["prompt"]) + 5
                ):
                    # chỉ thay khi prompt nguồn kết thúc bằng A lẻ hoặc ngắn hơn rõ
                    if item["prompt"].rstrip().endswith("A") or any(
                        not (opts[i] or "").strip() for i in range(4)
                    ):
                        item["prompt"] = pdf_q["prompt"]
                        changed = True

            # bỏ ghi chú "Thiếu phương án A..."
            note = item.get("note") or ""
            if "Thiếu phương án" in note:
                note = re.sub(r"\s*—\s*Thiếu phương án A trong nguồn DA340", "", note)
                note = re.sub(r"Thiếu phương án A trong nguồn DA340\s*—?\s*", "", note)
                item["note"] = note.strip(" —")
                changed = True

            if changed:
                pass

        # kiểm tra còn thiếu
        for i, key in enumerate("ABCD"):
            if not (item["options"][i] or "").strip():
                still_bad.append(f"{item['stt']}{key}")

    # Force fill ALL DA340 options from PDF when available (safer, complete)
    force_filled = 0
    for item in bank:
        if item.get("source") != "DA340":
            continue
        src_stt = item.get("source_stt")
        if src_stt not in pdf_qs:
            continue
        pdf_q = pdf_qs[src_stt]
        # chỉ overwrite option nếu PDF có đủ 4 phương án
        if all(pdf_q["options"][k] for k in "ABCD"):
            before = list(item["options"])
            item["options"] = [
                pdf_q["options"]["A"],
                pdf_q["options"]["B"],
                pdf_q["options"]["C"],
                pdf_q["options"]["D"],
            ]
            # sửa prompt nếu bị dính đuôi A
            if item["prompt"].rstrip().endswith("A") and pdf_q["prompt"]:
                item["prompt"] = pdf_q["prompt"]
            elif not item["prompt"] and pdf_q["prompt"]:
                item["prompt"] = pdf_q["prompt"]
            if before != item["options"]:
                force_filled += 1
            note = item.get("note") or ""
            if "Thiếu phương án" in note:
                note = re.sub(
                    r"\s*—\s*Thiếu phương án A trong nguồn DA340", "", note
                )
                item["note"] = note.strip(" —")

    still_bad = []
    for item in bank:
        for i, key in enumerate("ABCD"):
            if not (item["options"][i] or "").strip():
                still_bad.append(f"{item['stt']}{key}")

    print(f"Updated option sets from PDF: {force_filled} questions changed")
    if still_bad:
        print("STILL missing:", still_bad)
    else:
        print("OK: mọi câu 1–390 đủ A–D")

    # preview fixed ones
    for stt in (95, 294, 297, 298, 319, 321):
        q = bank[stt - 1]
        print(f"--- {stt} ---")
        print("Q:", q["prompt"][:100])
        for i, k in enumerate("ABCD"):
            print(f"  {k}. {q['options'][i][:90]}")

    OUT_JSON.write_text(json.dumps(bank, ensure_ascii=False, indent=2), encoding="utf-8")
    write_csv(bank, OUT_CSV)
    write_docx(bank, OUT_DOCX)
    print("Wrote:", OUT_DOCX)
    print("Wrote:", OUT_CSV)
    print("Wrote:", OUT_JSON)


if __name__ == "__main__":
    main()
