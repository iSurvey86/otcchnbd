# -*- coding: utf-8 -*-
"""Ghép DA340 + 50 câu TB 1952 → 390 câu dạng bảng chuẩn (Word + CSV/JSON cho DB)."""

from __future__ import annotations

import csv
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

DIR = Path(r"g:\My Drive\AIProject_Data\Thi CCHN\Dau thau")
FILE_50 = DIR / "Bo sung 50 cau hoi thi NVCM dau thau theo TB 1952 (lan 2) - bang.docx"
FILE_50_FALLBACK = DIR / (
    "Bo sung 50 cau hoi thi NVCM dau thau theo TB 1952 (lan 2) - khong watermark.docx"
)
FILE_340 = DIR / "DA340.docx"
OUT_DOCX = DIR / "NVCM-dau-thau-390-cau-bang.docx"
OUT_CSV = DIR / "NVCM-dau-thau-390-cau.csv"
OUT_JSON = DIR / "NVCM-dau-thau-390-cau.json"


def set_run_font(run, name: str = "Times New Roman", size: int = 11, bold: bool = False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_text(cell, text: str, *, size=11, bold=False, align=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_options_cell(cell, options: dict[str, str], *, size=11) -> None:
    cell.text = ""
    first = True
    for key in ("A", "B", "C", "D"):
        text = f"{key}. {options.get(key, '').strip()}"
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        set_run_font(run, size=size, bold=False)


def shade_header(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), "D9E2F3")
    shd.set(qn("w:val"), "clear")


def set_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "666666")
        borders.append(el)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def clean(s: str) -> str:
    s = re.sub(r"[ \t]+", " ", (s or "").replace("\xa0", " "))
    s = re.sub(r"\s+\n", "\n", s)
    return s.strip()


def format_answer_note(answer: str, note: str = "") -> str:
    answer = clean(answer).upper()
    note = clean(note)
    parts = []
    if answer in {"A", "B", "C", "D"}:
        parts.append(answer)
    if note:
        parts.append(note)
    return " — ".join(parts) if len(parts) > 1 else (parts[0] if parts else "")


def repair_options(prompt: str, options: dict[str, str]) -> tuple[str, dict[str, str], str]:
    """Sửa lỗi dính/thiếu phương án thường gặp trong DA340."""
    note_fix = ""
    opts = {k: clean(options.get(k, "")) for k in "ABCD"}
    prompt = clean(prompt)

    # B dính C: "... thầu C.Nhà thầu..."
    for key, nxt in (("A", "B"), ("B", "C"), ("C", "D")):
        if opts[key] and not opts[nxt]:
            m = re.search(rf"\s+{nxt}\.\s*", opts[key])
            if m:
                left = clean(opts[key][: m.start()])
                right = clean(opts[key][m.end() :])
                opts[key] = left
                opts[nxt] = right

    # Prompt dính chữ A ở cuối khi ô A trống: "...như thế nào?A"
    if not opts["A"]:
        m = re.search(r"([.?])\s*A\s*$", prompt)
        if m:
            prompt = clean(prompt[: m.end(1)])
        note_fix = "Thiếu phương án A trong nguồn DA340"

    return prompt, opts, note_fix


def load_340(path: Path) -> list[dict]:
    doc = Document(str(path))
    if not doc.tables:
        raise SystemExit(f"No table in {path}")
    table = doc.tables[0]
    out: list[dict] = []
    for row in table.rows[1:]:
        cells = [clean(c.text) for c in row.cells]
        if len(cells) < 6:
            continue
        stt_raw = cells[0]
        if not stt_raw.isdigit():
            continue
        prompt = cells[1]
        options = {"A": cells[2], "B": cells[3], "C": cells[4], "D": cells[5]}
        prompt, options, fix_note = repair_options(prompt, options)

        answer = cells[11] if len(cells) > 11 else ""
        answer_340 = cells[14] if len(cells) > 14 else ""
        evidence = cells[13] if len(cells) > 13 else ""
        status = cells[9] if len(cells) > 9 else ""

        def letter(s: str) -> str:
            s = clean(s).upper()
            return s[0] if s and s[0] in "ABCD" else ""

        lit = letter(answer_340) or letter(answer)
        note_parts: list[str] = []
        if lit:
            note_parts.append(lit)
            if evidence:
                note_parts.append(evidence)
        else:
            if status and status.lower() not in {"giữ nguyên", ""}:
                note_parts.append(f"Chưa có đáp án ({status})")
        if fix_note:
            note_parts.append(fix_note)

        out.append(
            {
                "source": "DA340",
                "source_stt": int(stt_raw),
                "prompt": prompt,
                "options": options,
                "answer": lit,
                "note": " — ".join(note_parts),
            }
        )
    return out


def load_50_table(path: Path) -> list[dict]:
    doc = Document(str(path))
    if not doc.tables:
        raise SystemExit(f"No table in {path}")
    table = doc.tables[0]
    out: list[dict] = []
    for row in table.rows[1:]:
        cells = [clean(c.text) for c in row.cells]
        if len(cells) < 3 or not cells[0].isdigit():
            continue
        opt_text = cells[2]
        options = {"A": "", "B": "", "C": "", "D": ""}
        # Tách A/B/C/D từ ô nhiều đoạn
        parts = re.split(r"(?m)^(?=[A-D]\.\s)", opt_text)
        # hoặc theo newline đã nối bằng \n trong cell.text
        chunks = re.findall(
            r"([A-D])\.\s*(.*?)(?=(?:[A-D]\.\s)|\Z)",
            opt_text,
            flags=re.S,
        )
        for key, body in chunks:
            options[key] = clean(body)
        answer_cell = cells[3] if len(cells) > 3 else ""
        m = re.match(r"^([A-Da-d])\b", answer_cell)
        letter = m.group(1).upper() if m else ""
        out.append(
            {
                "source": "TB1952-lan2",
                "source_stt": int(cells[0]),
                "prompt": cells[1],
                "options": options,
                "answer": letter,
                "note": clean(answer_cell) if not letter else (
                    answer_cell if len(clean(answer_cell)) > 1 else ""
                ),
            }
        )
    return out


def load_50_paras(path: Path) -> list[dict]:
    """Fallback: file dạng danh sách Câu n. / A. B. C. D."""
    doc = Document(str(path))
    lines = [clean(p.text) for p in doc.paragraphs if clean(p.text)]
    out: list[dict] = []
    current = None
    mode = None
    q_re = re.compile(r"^Câu\s+(\d+)\.\s*(.*)$", re.I)
    opt_re = re.compile(r"^([A-D])\s*\.\s*(.*)$")
    for line in lines:
        qm = q_re.match(line)
        if qm:
            current = {
                "source": "TB1952-lan2",
                "source_stt": int(qm.group(1)),
                "prompt": clean(qm.group(2)),
                "options": {"A": "", "B": "", "C": "", "D": ""},
                "answer": "",
                "note": "",
            }
            out.append(current)
            mode = "prompt"
            continue
        if current is None:
            continue
        om = opt_re.match(line)
        if om:
            mode = om.group(1)
            current["options"][mode] = clean(om.group(2))
            continue
        if mode == "prompt":
            current["prompt"] = clean(current["prompt"] + " " + line)
        elif mode in "ABCD":
            current["options"][mode] = clean(current["options"][mode] + " " + line)
    return out


def load_50() -> list[dict]:
    if FILE_50.exists():
        qs = load_50_table(FILE_50)
        if len(qs) == 50:
            return qs
    if FILE_50_FALLBACK.exists():
        qs = load_50_paras(FILE_50_FALLBACK)
        if len(qs) == 50:
            return qs
    raise SystemExit("Không đọc được đủ 50 câu TB 1952")


def write_docx(questions: list[dict], path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run(
        "Ngân hàng câu hỏi NVCM đấu thầu — 390 câu (340 + 50 bổ sung TB 1952 lần 2)"
    )
    set_run_font(r, size=13, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(10)
    r2 = sub.add_run(
        "STT 1–340: bộ DA340 (có đáp án/dẫn chứng nếu có). "
        "STT 341–390: bổ sung TB 1952/TB-QLĐT (chưa có đáp án — để trống)."
    )
    set_run_font(r2, size=10, bold=False)

    headers = (
        "STT",
        "Nội dung câu hỏi",
        "Phương án trả lời\n(Liệt kê đầy đủ A, B, C, D)",
        "Đáp án / Ghi chú",
    )
    table = doc.add_table(rows=1 + len(questions), cols=4)
    table.style = "Table Grid"
    table.autofit = False
    set_table_borders(table)
    widths = (Cm(1.2), Cm(8.3), Cm(13.2), Cm(4.0))
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w

    for i, h in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[i],
            h,
            size=11,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_header(table.rows[0].cells[i])

    for qi, q in enumerate(questions):
        row = table.rows[qi + 1]
        set_cell_text(
            row.cells[0],
            str(q["stt"]),
            size=11,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_cell_text(row.cells[1], q["prompt"], size=10)
        add_options_cell(row.cells[2], q["options"], size=10)
        note = q.get("note") or ""
        if not note and not q.get("answer"):
            note = ""  # để trống — chưa có đáp án
        set_cell_text(row.cells[3], note, size=10)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def write_csv(questions: list[dict], path: Path) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(
            [
                "stt",
                "prompt",
                "option_a",
                "option_b",
                "option_c",
                "option_d",
                "answer",
                "note",
                "source",
                "source_stt",
            ]
        )
        for q in questions:
            w.writerow(
                [
                    q["stt"],
                    q["prompt"],
                    q["options"]["A"],
                    q["options"]["B"],
                    q["options"]["C"],
                    q["options"]["D"],
                    q.get("answer", ""),
                    q.get("note", ""),
                    q.get("source", ""),
                    q.get("source_stt", ""),
                ]
            )


def write_json(questions: list[dict], path: Path) -> None:
    payload = []
    for q in questions:
        payload.append(
            {
                "stt": q["stt"],
                "prompt": q["prompt"],
                "options": [
                    q["options"]["A"],
                    q["options"]["B"],
                    q["options"]["C"],
                    q["options"]["D"],
                ],
                "answer": q.get("answer") or None,
                "note": q.get("note") or "",
                "source": q.get("source"),
                "source_stt": q.get("source_stt"),
            }
        )
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> None:
    import sys

    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

    q340 = load_340(FILE_340)
    q50 = load_50()
    print(f"Loaded DA340: {len(q340)}")
    print(f"Loaded TB1952: {len(q50)}")
    if len(q340) != 340:
        print(f"WARN: expected 340, got {len(q340)}")
    if len(q50) != 50:
        raise SystemExit(f"expected 50, got {len(q50)}")

    merged: list[dict] = []
    for q in q340:
        item = dict(q)
        item["stt"] = len(merged) + 1
        # chuẩn hóa note
        if item.get("answer") and item["answer"] in "ABCD":
            # note đã gồm đáp án + dẫn chứng
            if not item.get("note"):
                item["note"] = item["answer"]
        else:
            item["answer"] = ""
            item["note"] = item.get("note") or ""
        merged.append(item)

    for q in q50:
        item = dict(q)
        item["stt"] = len(merged) + 1
        # chưa có đáp án → để trống
        if not item.get("answer"):
            item["note"] = ""
        merged.append(item)

    assert len(merged) == 390, len(merged)

    # kiểm tra thiếu option
    bad = []
    for q in merged:
        for k in "ABCD":
            if not q["options"].get(k):
                bad.append(f"{q['stt']}{k}")
    with_ans = sum(1 for q in merged if q.get("answer"))
    print(f"Total: {len(merged)}; có đáp án: {with_ans}; thiếu đáp án: {len(merged) - with_ans}")
    if bad:
        print("WARN missing options:", bad[:20], "..." if len(bad) > 20 else "")

    write_docx(merged, OUT_DOCX)
    write_csv(merged, OUT_CSV)
    write_json(merged, OUT_JSON)
    print("Wrote:")
    print(" ", OUT_DOCX)
    print(" ", OUT_CSV)
    print(" ", OUT_JSON)


if __name__ == "__main__":
    main()
