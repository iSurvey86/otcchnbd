# -*- coding: utf-8 -*-
"""PDF TB 1952 (50 câu NVCM đấu thầu) → Word, bỏ watermark ảnh nền."""

from __future__ import annotations

import re
from pathlib import Path

import fitz
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PDF = Path(
    r"g:\My Drive\AIProject_Data\Thi CCHN\Dau thau"
    r"\Bổ sung 50 câu hỏi thi NVCM đấu thầu theo TB 1952 (lần 2).pdf"
)
OUT = Path(
    r"g:\My Drive\AIProject_Data\Thi CCHN\Dau thau"
    r"\Bo-sung-50-cau-hoi-NVCM-dau-thau-TB-1952-lan-2.docx"
)

SKIP = {
    "STT",
    "Nội dung",
    "Phương án trả lời",
}
OPTION_START = re.compile(r"^([A-D])\s*\.\s*(.*)$", re.S)
STT_ONLY = re.compile(r"^(\d{1,2})$")
TITLE_START = re.compile(r"^Bổ sung 50 câu hỏi", re.I)


def set_run_font(run, name: str = "Times New Roman", size: int = 12, bold: bool = False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc, text: str, *, size=12, bold=False, space_after=6, space_before=0, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def normalize_spaces(s: str) -> str:
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r" *\n *", " ", s)
    s = re.sub(r"\s+([,.;:!?])", r"\1", s)
    # sửa lỗi dính chữ kiểu "A.Không"
    s = re.sub(r"^([A-D])\.(\S)", r"\1. \2", s)
    return s.strip()


def extract_text_lines(pdf: fitz.Document) -> list[str]:
    """Lấy text theo đúng thứ tự đọc; bỏ ảnh watermark (không nằm trong text)."""
    lines: list[str] = []
    for page in pdf:
        raw = page.get_text("text")
        for line in raw.splitlines():
            t = line.strip()
            if t:
                lines.append(t)
    return lines


def parse_questions(lines: list[str]) -> tuple[str, list[dict]]:
    title_parts: list[str] = []
    questions: list[dict] = []
    current: dict | None = None
    mode: str | None = None  # 'prompt' | 'A'|'B'|'C'|'D'

    def append_text(bucket: str, text: str) -> None:
        assert current is not None
        if bucket == "prompt":
            prev = current["prompt"]
            joiner = "" if not prev or prev.endswith((" ", "-", "/")) else " "
            current["prompt"] = prev + joiner + text
        else:
            prev = current["options"][bucket]
            joiner = "" if not prev or prev.endswith((" ", "-", "/")) else " "
            current["options"][bucket] = prev + joiner + text

    for text in lines:
        if text in SKIP:
            continue
        if TITLE_START.search(text) or text.startswith("phục vụ ôn luyện"):
            title_parts.append(text)
            continue

        stt_m = STT_ONLY.match(text)
        if stt_m:
            n = int(stt_m.group(1))
            if 1 <= n <= 50:
                current = {
                    "stt": n,
                    "prompt": "",
                    "options": {"A": "", "B": "", "C": "", "D": ""},
                }
                questions.append(current)
                mode = "prompt"
                continue

        if current is None:
            continue

        opt_m = OPTION_START.match(text)
        if opt_m:
            mode = opt_m.group(1)
            rest = opt_m.group(2).strip()
            if rest:
                append_text(mode, rest)
            continue

        if mode:
            append_text(mode, text)

    title = " ".join(title_parts).strip()
    return title, questions


def set_cell_text(cell, text: str, *, size=11, bold=False, align=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_options_cell(cell, options: dict[str, str], *, size=11) -> None:
    cell.text = ""
    first = True
    for key in ("A", "B", "C", "D"):
        text = f"{key}. {normalize_spaces(options.get(key, ''))}"
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        set_run_font(run, size=size, bold=False)


def shade_header(cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        from docx.oxml import OxmlElement

        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), "D9E2F3")
    shd.set(qn("w:val"), "clear")


def set_table_borders(table) -> None:
    from docx.oxml import OxmlElement

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "666666")
        borders.append(el)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def build_docx(title: str, questions: list[dict], out: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    # Ngang A4 — bảng 4 cột dễ đọc hơn
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    add_para(
        doc,
        "Bổ sung 50 câu hỏi trắc nghiệm — Nghiệp vụ chuyên môn về đấu thầu",
        size=13,
        bold=True,
        space_after=2,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        "Theo Thông báo số 1952/TB-QLĐT ngày 19/9/2025 của Cục Quản lý đấu thầu "
        "(ôn luyện từ 30/01/2026) — đã loại bỏ watermark",
        size=10,
        bold=False,
        space_after=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    headers = (
        "STT",
        "Nội dung câu hỏi",
        "Phương án trả lời\n(Liệt kê đầy đủ A, B, C, D)",
        "Đáp án / Ghi chú",
    )
    table = doc.add_table(rows=1 + len(questions), cols=4)
    table.style = "Table Grid"
    table.autofit = False
    set_table_borders(table)

    # Độ rộng cột (tổng ~26.7 cm usable)
    widths = (Cm(1.2), Cm(8.5), Cm(13.5), Cm(3.5))
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w

    header_row = table.rows[0]
    for i, h in enumerate(headers):
        set_cell_text(
            header_row.cells[i],
            h,
            size=11,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_header(header_row.cells[i])

    missing: list[str] = []
    for qi, q in enumerate(questions):
        row = table.rows[qi + 1]
        stt = q["stt"]
        prompt = normalize_spaces(q["prompt"])
        if not prompt:
            missing.append(f"Q{stt}")
        for key in ("A", "B", "C", "D"):
            if not normalize_spaces(q["options"].get(key, "")):
                missing.append(f"Q{stt}{key}")

        set_cell_text(
            row.cells[0],
            str(stt),
            size=11,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_cell_text(row.cells[1], prompt, size=11, bold=False)
        add_options_cell(row.cells[2], q["options"], size=11)
        set_cell_text(row.cells[3], "", size=11)  # PDF không có đáp án

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"Wrote: {out}")
    print(f"Questions: {len(questions)}")
    if missing:
        print("WARN missing:", missing)
    else:
        print("OK: đủ 50 câu × A–D (bảng 4 cột)")


def main() -> None:
    import shutil
    import sys

    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

    pdf = fitz.open(str(PDF))
    lines = extract_text_lines(pdf)
    title, questions = parse_questions(lines)
    by_stt = {q["stt"]: q for q in questions if 1 <= q["stt"] <= 50}
    ordered = [by_stt[i] for i in range(1, 51) if i in by_stt]
    if len(ordered) != 50:
        print("FOUND:", sorted(by_stt))
        raise SystemExit(f"Expected 50, got {len(ordered)}")

    build_docx(title, ordered, OUT)
    # Bản tên dễ đọc
    alt = OUT.parent / (
        "Bo sung 50 cau hoi thi NVCM dau thau theo TB 1952 (lan 2) - bang.docx"
    )
    shutil.copy2(OUT, alt)
    print(f"Also: {alt}")


if __name__ == "__main__":
    main()
